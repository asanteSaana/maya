"""Shared Word-document toolkit for the capstone deliverables.

Every document is produced from the same helpers so headings, tables, figures
and code listings look identical across the submission set.
"""

import os

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DIAGRAMS = os.path.join(ROOT, "diagrams")
OUT = ROOT
os.makedirs(OUT, exist_ok=True)

GREEN = RGBColor(0x27, 0x4C, 0x3B)
GREEN_MID = RGBColor(0x4F, 0x7E, 0x63)
INK = RGBColor(0x23, 0x30, 0x20)
GREY = RGBColor(0x6B, 0x72, 0x6C)
RED = RGBColor(0xA8, 0x3A, 0x2E)

BODY_FONT = "Calibri"
MONO_FONT = "Consolas"

# Filled in from facts.py; kept here so a single edit updates every document.
PLACEHOLDER_STYLE = "[[ ]]"


def _shade(cell, hex_colour):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_colour)
    tc_pr.append(shd)


def _para_shade(paragraph, hex_colour):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_colour)
    p_pr.append(shd)


def new_document(subtitle=None):
    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    for attr, value in (
        ("top_margin", 2.2), ("bottom_margin", 2.2),
        ("left_margin", 2.4), ("right_margin", 2.2),
    ):
        setattr(section, attr, Cm(value))

    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.18

    for level, size in ((1, 16), (2, 13), (3, 11.5)):
        style = doc.styles[f"Heading {level}"]
        style.font.name = BODY_FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = GREEN if level < 3 else GREEN_MID
        style.paragraph_format.space_before = Pt(14 if level == 1 else 10)
        style.paragraph_format.space_after = Pt(5)
        style.paragraph_format.keep_with_next = True

    if subtitle:
        add_footer(doc, subtitle)
    return doc


def add_footer(doc, text):
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{text}    |    ")
    run.font.size = Pt(8)
    run.font.color.rgb = GREY

    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    p._p.append(fld)


def title_page(doc, title, doc_kind, facts):
    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(facts.COURSE)
    run.font.size = Pt(12)
    run.font.color.rgb = GREY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = GREEN

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(doc_kind)
    run.font.size = Pt(14)
    run.font.color.rgb = GREEN_MID

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(facts.TAGLINE)
    run.font.size = Pt(11)
    run.italic = True
    run.font.color.rgb = GREY

    doc.add_paragraph()
    doc.add_paragraph()

    rows = [
        ("Group Number / Name", facts.GROUP_NAME),
        ("Project Title", facts.PROJECT_TITLE),
        ("Document", doc_kind),
        ("Version", facts.VERSION),
        ("Date", facts.DATE),
        ("Live application", facts.LIVE_URL),
        ("Source repository", facts.REPO_URL),
    ]
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].width = Cm(5.2)
        cells[1].width = Cm(10.4)
        run = cells[0].paragraphs[0].add_run(label)
        run.bold = True
        run.font.size = Pt(10)
        _shade(cells[0], "EFF4F0")
        add_value(cells[1].paragraphs[0], value, facts)

    doc.add_paragraph()
    members_table(doc, facts)
    doc.add_page_break()


def add_value(paragraph, value, facts):
    """Writes a value, flagging anything still awaiting the group's input."""
    run = paragraph.add_run(value)
    run.font.size = Pt(10)
    if facts.is_placeholder(value):
        run.font.color.rgb = RED
        run.bold = True
        highlight(run)


def highlight(run):
    r_pr = run._r.get_or_add_rPr()
    hl = OxmlElement("w:highlight")
    hl.set(qn("w:val"), "yellow")
    r_pr.append(hl)


def members_table(doc, facts):
    p = doc.add_paragraph()
    run = p.add_run("Group Members")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = GREEN

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ("No.", "Group Member", "Student ID", "Major Contribution")
    widths = (Cm(1.2), Cm(4.6), Cm(3.2), Cm(6.8))
    for cell, text, width in zip(table.rows[0].cells, headers, widths):
        cell.width = width
        run = cell.paragraphs[0].add_run(text)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _shade(cell, "274C3B")

    for index, member in enumerate(facts.MEMBERS, start=1):
        cells = table.add_row().cells
        cells[0].paragraphs[0].add_run(str(index)).font.size = Pt(10)
        for cell, value, width in zip(cells[1:], member, widths[1:]):
            cell.width = width
            add_value(cell.paragraphs[0], value, facts)


def toc(doc):
    doc.add_heading("Table of Contents", level=1)
    p = doc.add_paragraph()
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Right-click here and choose “Update Field” to build the contents list."
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    for element in (fld_begin, instr, fld_sep, placeholder, fld_end):
        run._r.append(element)
    doc.add_page_break()


def h1(doc, text):
    return doc.add_heading(text, level=1)


def h2(doc, text):
    return doc.add_heading(text, level=2)


def h3(doc, text):
    return doc.add_heading(text, level=3)


def para(doc, text, italic=False, size=11, align=None, colour=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = italic
    run.font.size = Pt(size)
    if colour is not None:
        run.font.color.rgb = colour
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "justify":
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def bullets(doc, items, style="List Bullet"):
    for item in items:
        p = doc.add_paragraph(style=style)
        if isinstance(item, tuple):
            lead, rest = item
            run = p.add_run(lead)
            run.bold = True
            p.add_run(rest)
        else:
            p.add_run(item)
        p.paragraph_format.space_after = Pt(3)


def numbered(doc, items):
    bullets(doc, items, style="List Number")


def table(doc, headers, rows, widths=None, font_size=9.5, first_col_bold=False):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, (cell, text) in enumerate(zip(t.rows[0].cells, headers)):
        if widths:
            cell.width = Cm(widths[i])
        run = cell.paragraphs[0].add_run(text)
        run.bold = True
        run.font.size = Pt(font_size)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _shade(cell, "274C3B")

    for r, row in enumerate(rows):
        cells = t.add_row().cells
        for i, (cell, value) in enumerate(zip(cells, row)):
            if widths:
                cell.width = Cm(widths[i])
            run = cell.paragraphs[0].add_run(str(value))
            run.font.size = Pt(font_size)
            if first_col_bold and i == 0:
                run.bold = True
            cell.paragraphs[0].paragraph_format.space_after = Pt(2)
        if r % 2 == 1:
            for cell in cells:
                _shade(cell, "F4F7F5")
    doc.add_paragraph()
    return t


def figure(doc, filename, caption, width_cm=15.5):
    path = os.path.join(DIAGRAMS, filename)
    if not os.path.exists(path):
        para(doc, f"[missing diagram: {filename}]", italic=True, colour=RED)
        return
    doc.add_picture(path, width=Cm(width_cm))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(caption)
    run.font.size = Pt(9)
    run.italic = True
    run.font.color.rgb = GREY


def code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.name = MONO_FONT
    run.font.size = Pt(8.5)
    r_pr = run._r.get_or_add_rPr()
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), MONO_FONT)
    fonts.set(qn("w:hAnsi"), MONO_FONT)
    r_pr.append(fonts)
    _para_shade(p, "F4F7F5")
    return p


def callout(doc, title, text, colour="FFF6E5"):
    p = doc.add_paragraph()
    run = p.add_run(f"{title}  ")
    run.bold = True
    run.font.size = Pt(10)
    run2 = p.add_run(text)
    run2.font.size = Pt(10)
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(8)
    _para_shade(p, colour)
    return p


def page_break(doc):
    doc.add_page_break()


def save(doc, filename):
    path = os.path.join(OUT, filename)
    doc.save(path)
    return path
